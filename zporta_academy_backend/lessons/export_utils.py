# lessons/export_utils.py
"""
Utilities for exporting lesson content to various formats (PDF, DOCX)
Uses WeasyPrint for high-quality PDF generation with full HTML/CSS support
"""

from io import BytesIO
from django.conf import settings
from bs4 import BeautifulSoup
import re

def clean_html_for_export(html_content):
    """Clean HTML for export while preserving structure"""
    if not html_content:
        return ""
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove script tags only (keep styles for rendering)
    for tag in soup.find_all('script'):
        tag.decompose()
    
    # Remove contenteditable attributes
    for tag in soup.find_all(attrs={"contenteditable": True}):
        del tag['contenteditable']
    
    return str(soup)


def build_export_html(lesson):
    """Build complete HTML document for export with all styling"""
    clean_content = clean_html_for_export(lesson.content)
    accent_color = lesson.accent_color or '#222E3B'
    custom_css = lesson.custom_css or ''
    
    html_template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{lesson.title}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 100%;
        }}
        
        .lesson-header {{
            text-align: center;
            margin-bottom: 2rem;
            padding-bottom: 1rem;
            border-bottom: 3px solid {accent_color};
        }}
        
        .lesson-header h1 {{
            color: {accent_color};
            font-size: 2rem;
            margin-bottom: 0.5rem;
        }}
        
        .lesson-meta {{
            color: #666;
            font-size: 0.9rem;
            margin-bottom: 0.5rem;
        }}
        
        .lesson-content {{
            margin-top: 2rem;
        }}
        
        .lesson-content img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1rem auto;
            border-radius: 8px;
        }}
        
        .lesson-content video {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 1rem auto;
        }}
        
        .lesson-content h1, .lesson-content h2, .lesson-content h3 {{
            color: {accent_color};
            margin-top: 1.5rem;
            margin-bottom: 0.75rem;
        }}
        
        .lesson-content p {{
            margin-bottom: 1rem;
        }}
        
        .lesson-content ul, .lesson-content ol {{
            margin-bottom: 1rem;
            padding-left: 2rem;
        }}
        
        .lesson-content li {{
            margin-bottom: 0.5rem;
        }}
        
        .lesson-content table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 1rem;
        }}
        
        .lesson-content table th,
        .lesson-content table td {{
            border: 1px solid #ddd;
            padding: 0.75rem;
            text-align: left;
        }}
        
        .lesson-content table th {{
            background-color: {accent_color};
            color: white;
            font-weight: 600;
        }}
        
        .lesson-content blockquote {{
            border-left: 4px solid {accent_color};
            padding-left: 1rem;
            margin: 1rem 0;
            font-style: italic;
            color: #555;
        }}
        
        .lesson-content code {{
            background-color: #f4f4f4;
            padding: 0.2rem 0.4rem;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 0.9em;
        }}
        
        .lesson-content pre {{
            background-color: #f4f4f4;
            padding: 1rem;
            border-radius: 5px;
            overflow-x: auto;
            margin-bottom: 1rem;
        }}
        
        .lesson-content pre code {{
            background-color: transparent;
            padding: 0;
        }}
        
        /* Zporta custom components */
        .zporta-columns {{
            display: grid;
            gap: 1.5rem;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            margin: 1rem 0;
        }}
        
        .zporta-column {{
            min-width: 0;
        }}
        
        .zporta-button {{
            display: inline-block;
            padding: 0.6rem 1.1rem;
            border-radius: 8px;
            text-decoration: none;
            font-weight: 600;
            margin: 0.5rem 0.5rem 0.5rem 0;
        }}
        
        .zporta-btn--primary {{
            background-color: {accent_color};
            color: white;
            border: 1px solid {accent_color};
        }}
        
        .zporta-btn--secondary {{
            background-color: white;
            color: {accent_color};
            border: 1px solid {accent_color};
        }}
        
        .zporta-acc-item {{
            border: 1px solid #e2e8f0;
            border-radius: 8px;
            margin-bottom: 0.75rem;
            overflow: hidden;
        }}
        
        .zporta-acc-title {{
            padding: 0.75rem 1rem;
            background-color: #f8fafc;
            font-weight: 600;
            color: {accent_color};
        }}
        
        .zporta-acc-panel {{
            padding: 1rem;
            border-top: 1px solid #e2e8f0;
        }}
        
        {custom_css}
    </style>
</head>
<body>
    <div class="lesson-header">
        <h1>{lesson.title}</h1>
        <div class="lesson-meta">
            <span>Created by: {lesson.created_by.username}</span> | 
            <span>Date: {lesson.created_at.strftime('%B %d, %Y')}</span>
        </div>
        {f'<div class="lesson-meta">Subject: {lesson.subject.name}</div>' if lesson.subject else ''}
        {f'<div class="lesson-meta">Course: {lesson.course.title}</div>' if lesson.course else ''}
    </div>
    
    <div class="lesson-content">
        {clean_content}
    </div>
    
    {f'<div style="margin-top: 2rem; padding-top: 1rem; border-top: 1px solid #ddd;"><strong>Video:</strong> <a href="{lesson.video_url}">{lesson.video_url}</a></div>' if lesson.video_url else ''}
</body>
</html>
    """
    
    return html_template.strip()


def generate_lesson_pdf(lesson):
    """
    Generate a PDF from lesson content with full HTML/CSS styling
    Uses WeasyPrint for high-quality rendering
    """
    try:
        from weasyprint import HTML, CSS
    except ImportError:
        # Fallback to simple reportlab if WeasyPrint not available
        return generate_lesson_pdf_simple(lesson)
    
    try:
        # Build complete HTML with styling
        html_content = build_export_html(lesson)
        
        # Generate PDF
        pdf_file = BytesIO()
        HTML(string=html_content).write_pdf(pdf_file)
        pdf_data = pdf_file.getvalue()
        pdf_file.close()
        
        return pdf_data, None
    except Exception as e:
        return None, f"PDF generation error: {str(e)}"


def generate_lesson_pdf_simple(lesson):
    """
    Fallback: Generate a simple PDF from lesson content (plain text)
    Requires: pip install reportlab
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.pdfgen import canvas
    except ImportError:
        return None, "ReportLab not installed. Run: pip install reportlab"
    
    buffer = BytesIO()
    
    # Create PDF
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
    
    # Title
    title_style = styles['Heading1']
    title_style.alignment = TA_CENTER
    elements.append(Paragraph(lesson.title, title_style))
    elements.append(Spacer(1, 12))
    
    # Metadata
    meta_text = f"Created by: {lesson.created_by.username}<br/>"
    meta_text += f"Created: {lesson.created_at.strftime('%Y-%m-%d')}<br/>"
    if lesson.subject:
        meta_text += f"Subject: {lesson.subject.name}<br/>"
    if lesson.course:
        meta_text += f"Course: {lesson.course.title}<br/>"
    
    elements.append(Paragraph(meta_text, styles['Normal']))
    elements.append(Spacer(1, 24))
    
    # Content
    clean_content = clean_html_for_export(lesson.content)
    
    # Convert HTML to plain text for PDF (simple approach)
    soup = BeautifulSoup(clean_content, 'html.parser')
    text_content = soup.get_text()
    
    # Split into paragraphs and add to PDF
    for para in text_content.split('\n'):
        if para.strip():
            elements.append(Paragraph(para, styles['Normal']))
            elements.append(Spacer(1, 12))
    
    # Video URL if exists
    if lesson.video_url:
        elements.append(Spacer(1, 24))
        elements.append(Paragraph(f"Video: {lesson.video_url}", styles['Normal']))
    
    # Build PDF
    doc.build(elements)
    
    # Get PDF data
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data, None


def generate_lesson_docx(lesson):
    """
    Generate a DOCX from lesson content with better HTML parsing
    Requires: pip install python-docx htmldocx
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None, "python-docx not installed. Run: pip install python-docx"
    
    try:
        from htmldocx import HtmlToDocx
        use_html_parser = True
    except ImportError:
        use_html_parser = False
    
    buffer = BytesIO()
    document = Document()
    
    # Title
    title = document.add_heading(lesson.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    document.add_paragraph(f"Created by: {lesson.created_by.username}")
    document.add_paragraph(f"Created: {lesson.created_at.strftime('%B %d, %Y')}")
    if lesson.subject:
        document.add_paragraph(f"Subject: {lesson.subject.name}")
    if lesson.course:
        document.add_paragraph(f"Course: {lesson.course.title}")
    
    document.add_paragraph()  # Empty line
    
    # Content
    clean_content = clean_html_for_export(lesson.content)
    
    if use_html_parser:
        # Use HtmlToDocx for better HTML rendering
        parser = HtmlToDocx()
        parser.add_html_to_document(clean_content, document)
    else:
        # Fallback to simple parsing
        soup = BeautifulSoup(clean_content, 'html.parser')
        
        # Extract text with basic formatting
        for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li', 'ul', 'ol']):
            if element.name == 'h1':
                document.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                document.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                document.add_heading(element.get_text(), level=3)
            elif element.name == 'p':
                text = element.get_text().strip()
                if text:
                    document.add_paragraph(text)
            elif element.name == 'li':
                text = element.get_text().strip()
                if text:
                    document.add_paragraph(text, style='List Bullet')
    
    # Video URL if exists
    if lesson.video_url:
        document.add_paragraph()
        p = document.add_paragraph(f"Video: {lesson.video_url}")
        p.runs[0].font.color.rgb = RGBColor(0, 0, 255)
    
    # Save to buffer
    document.save(buffer)
    
    # Get DOCX data
    docx_data = buffer.getvalue()
    buffer.close()
    
    return docx_data, None
